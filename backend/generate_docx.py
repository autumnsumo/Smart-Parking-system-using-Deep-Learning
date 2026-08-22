import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('Review-I: Seminar Presentation Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Project Title: Automated Smart Parking System using YOLOv8 and LPRNet', level=1)

# 1
doc.add_heading('1. Introduction and Motivation', level=2)
doc.add_paragraph('With the rapid increase in urbanization and vehicle ownership, finding parking spaces in metropolitan areas has become a significant challenge. This leads to traffic congestion, increased carbon emissions, and wasted time. Traditional parking management systems rely heavily on manual ticketing and human oversight, which are inefficient, error-prone, and labor-intensive.')
p = doc.add_paragraph()
p.add_run('Motivation: ').bold = True
p.add_run('The motivation behind this project is to automate the parking management process by leveraging cutting-edge deep learning technologies. By integrating an Automatic Number Plate Recognition (ANPR) system with a real-time web dashboard, parking administrators can seamlessly monitor occupancy, automate entry/exit protocols, and eliminate physical ticketing, resulting in a frictionless experience for both users and operators.')

# 2
doc.add_heading('2. Problem Statement', level=2)
doc.add_paragraph('Existing ANPR-based parking systems frequently struggle with low-resolution, noisy, or damaged license plates. General-purpose OCR engines (like Tesseract or EasyOCR) are not specifically tailored for the typography, spacing, and multi-line formats of Indian regional transport (RTO) plates, leading to hallucinated characters and high error rates. Furthermore, traditional detection methods (like Haar Cascades) are highly susceptible to varying lighting and weather conditions. There is a critical need for an end-to-end, highly accurate, and real-time ALPR pipeline optimized specifically for the Indian context.')

# 3
doc.add_heading('3. Objectives', level=2)
doc.add_paragraph('1. Robust Vehicle/Plate Detection: To implement a YOLOv8 object detection model capable of precisely localizing license plates in real-time under diverse environmental conditions.')
doc.add_paragraph('2. High-Accuracy Recognition: To replace general-purpose OCR with a specialized License Plate Recognition Network (LPRNet) fine-tuned on Indian plates to achieve near 100% accuracy.')
doc.add_paragraph('3. Automated Parking Management: To develop a dynamic region-of-interest (ROI) mapping tool that automatically tracks slot occupancy based on camera feeds.')
doc.add_paragraph('4. End-to-End Integration: To build a full-stack system featuring a FastAPI backend and a React frontend dashboard for seamless administration, real-time alerts, and automated billing.')

# 4
doc.add_heading('4. Literature Survey (Minimum 10 Recent IEEE/SCI Papers)', level=2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'S.No.'
hdr_cells[1].text = 'Paper Title (Year)'
hdr_cells[2].text = 'Core Technologies'
hdr_cells[3].text = 'Key Contribution'

papers = [
    ('1', 'Edge-Optimized Automatic Number Plate Recognition for IoT-Based Smart Parking (2026)', 'YOLOv8, PaddleOCR, Edge AI', 'Deployed YOLOv8 on Raspberry Pi for resource-constrained, sensor-triggered parking environments.'),
    ('2', 'Real-Time Vehicle Number Plate Recognition and Smart Parking Allocation (2026)', 'YOLOv8, CNN, OCR', 'Proposed automated dynamic slot allocation algorithms integrated with YOLOv8 plate extraction.'),
    ('3', 'CSCM-YOLOv8 and CSM-LPRNet for Unconstrained License Plate Recognition (2026)', 'YOLOv8, LPRNet, Attention', 'Introduced attention mechanisms to drastically improve recognition in rain, snow, and low-light scenarios.'),
    ('4', 'Attention-Based LPR-CBAM-Net for Robust Vehicle Identification (2026)', 'YOLOv5/8, LPRNet, CBAM', 'Integrated Convolutional Block Attention Modules to recalibrate feature channels for faster inference.'),
    ('5', 'Smart Parking System Using YOLOv8-Based Recognition and Automated Billing (2025)', 'YOLOv8, Dynamic Billing', 'Developed a complete ecosystem tracking vehicle dwell time for automated, dynamic fee calculation.'),
    ('6', 'Optimized YOLOv8 for Automatic License Plate Recognition on Edge Devices (2025)', 'YOLOv8-s, Edge Deployment', 'Achieved 99.3% mAP on benchmark ALPR datasets by optimizing the YOLOv8-small architecture.'),
    ('7', 'End-to-End Indian License Plate Recognition using Deep Learning (2025)', 'LPRNet, Spatial Transformer', 'Specifically curated an Indian dataset to train an LPRNet variant, overcoming regional font discrepancies.'),
    ('8', 'IoT-Enabled Smart Parking Management System using Machine Vision (2024)', 'OpenCV, YOLOv8, IoT Sensors', 'Combined visual AI with physical IR sensors to provide a highly redundant occupancy detection mechanism.'),
    ('9', 'Automatic Number Plate and Speed Detection using YOLO and CNN (IEEE I2CT, 2024)', 'YOLO, Custom CNN', 'Highlighted the failure modes of standard OCR and proposed custom CNN architectures for character extraction.'),
    ('10', 'Enhanced YOLOv8-Based System for Complex Environmental ANPR (MDPI, 2024)', 'YOLOv8s, Image Pre-processing', 'Explored the impact of CLAHE and luminosity normalization on detection confidence in poor lighting.')
]

for s, title, tech, cont in papers:
    row_cells = table.add_row().cells
    row_cells[0].text = s
    row_cells[1].text = title
    row_cells[2].text = tech
    row_cells[3].text = cont

# 5
doc.add_heading('5. Research Gap Identification', level=2)
doc.add_paragraph('1. Regional Incompatibility: Most pre-trained LPRNet and YOLO models are trained on Chinese (CCPD) or American datasets, resulting in poor accuracy when deployed on multi-line Indian license plates.')
doc.add_paragraph('2. System Latency vs. Accuracy: High-accuracy models often require heavy compute, making them unsuitable for real-time edge processing at parking booms.')
doc.add_paragraph('3. Lack of Integrated Workflows: Many papers focus solely on the AI recognition metric (mAP/Accuracy) but fail to integrate the model into a functional, user-facing Smart Parking ecosystem with slot management and billing.')

# 6
doc.add_heading('6. Proposed Methodology', level=2)
doc.add_paragraph('The proposed system addresses the research gaps by implementing a highly optimized two-stage pipeline explicitly tuned for the Indian context:')
doc.add_paragraph('Stage 1 (Detection): The video feed or captured image is passed through a YOLOv8 model. YOLOv8 outputs precise bounding box coordinates for the license plate, effectively filtering out background noise.', style='List Bullet')
doc.add_paragraph('Stage 2 (Recognition): The bounding box region is cropped, normalized, and resized to 94x24 pixels. This Region of Interest (ROI) is passed into LPRNet, an end-to-end convolutional neural network. LPRNet processes the entire plate holistically (without slicing individual characters) and uses a Connectionist Temporal Classification (CTC) greedy decoder to output the final string (e.g., MH20EE0943).', style='List Bullet')
doc.add_paragraph('Stage 3 (Management): The extracted plate is logged into an SQLite database via a FastAPI backend. The React frontend consumes this data via WebSockets/REST, updating dynamic UI elements like "Parking Full" statuses and individual slot occupancies.', style='List Bullet')

# 7
doc.add_heading('7. Algorithm / Flowchart', level=2)
doc.add_paragraph('1. Vehicle Arrives at Entry -> Camera Captures Image')
doc.add_paragraph('2. YOLOv8 Detection:')
doc.add_paragraph('   a) If Plate Detected -> Crop Bounding Box')
doc.add_paragraph('   b) If No Plate -> Trigger Alert / Retry')
doc.add_paragraph('3. Image Normalization & Resize to 94x24')
doc.add_paragraph('4. LPRNet CNN Feature Extraction')
doc.add_paragraph('5. CTC Greedy Decoder -> Extracted Plate String')
doc.add_paragraph('6. Check Database:')
doc.add_paragraph('   a) If New Vehicle -> Assign Empty Slot & Log Entry Time')
doc.add_paragraph('   b) If Existing Vehicle -> Log Exit Time & Calculate Bill')
doc.add_paragraph('7. Update React Dashboard UI')

# 8
doc.add_heading('8. System Architecture', level=2)
doc.add_paragraph('The architecture follows a modern, decoupled Client-Server model:')
doc.add_paragraph('Presentation Layer (Frontend): Built with React.js, Vite, and TailwindCSS. Handles ROI configuration, slot monitoring, entry/exit logs, and real-time toast notifications.', style='List Bullet')
doc.add_paragraph('Application Layer (Backend): Built with FastAPI (Python). Hosts the RESTful API endpoints and WebSocket servers for real-time bidirectional communication.', style='List Bullet')
doc.add_paragraph('AI Layer (Inference): YOLOv8 (Ultralytics) for detection and PyTorch LPRNet for recognition. Models are loaded directly into GPU VRAM on server startup for zero-latency inference.', style='List Bullet')
doc.add_paragraph('Data Layer (Database): SQLite/SQLAlchemy for persistent storage of vehicle logs, assigned slots, and ROI coordinates.', style='List Bullet')

# 9
doc.add_heading('9. Software & Hardware Requirements', level=2)
doc.add_paragraph('Hardware Requirements:', style='Heading 3')
doc.add_paragraph('- Processor: Intel Core i5 / AMD Ryzen 5 (Minimum)')
doc.add_paragraph('- RAM: 8 GB (16 GB Recommended)')
doc.add_paragraph('- GPU: NVIDIA GPU with CUDA support (e.g., GTX 1650 or higher) for real-time inference.')
doc.add_paragraph('- Camera: Standard 1080p IP Camera or Web Camera.')
doc.add_paragraph('Software Requirements:', style='Heading 3')
doc.add_paragraph('- OS: Windows 10/11 or Ubuntu 20.04+')
doc.add_paragraph('- Environment: Python 3.9+, Node.js 18+')
doc.add_paragraph('- Deep Learning Frameworks: PyTorch, Ultralytics (YOLO)')
doc.add_paragraph('- Database: SQLite3')

# 10
doc.add_heading('10. Tools & Development Environment', level=2)
doc.add_paragraph('- IDE: Visual Studio Code')
doc.add_paragraph('- Version Control: Git & GitHub')
doc.add_paragraph('- Backend Framework: FastAPI, Uvicorn, SQLAlchemy')
doc.add_paragraph('- Frontend Framework: React, Vite, Tailwind CSS, Lucide Icons')
doc.add_paragraph('- Computer Vision Libraries: OpenCV, NumPy')

# 11
doc.add_heading('11. Dataset Description (if applicable)', level=2)
doc.add_paragraph('The models were trained/fine-tuned on custom and publicly available Indian vehicle datasets:')
doc.add_paragraph('Detection Dataset: Thousands of annotated images of Indian cars, bikes, and trucks under various lighting conditions, annotated with YOLO bounding box formats.', style='List Bullet')
doc.add_paragraph('Recognition Dataset (Synthetic & Real): A combination of the Indian_LPR dataset and synthetically generated white-background/black-text images matching RTO fonts to train the LPRNet CTC decoder.', style='List Bullet')

# 12
doc.add_heading('12. Module Description', level=2)
doc.add_paragraph('1. ROI Configuration Module: Allows administrators to upload a static image of the parking lot and draw bounding polygons over valid parking spaces.')
doc.add_paragraph('2. Occupancy Detection Module: Periodically checks the defined ROI polygons to determine if a slot is "Vacant" or "Occupied".')
doc.add_paragraph('3. ANPR Entry/Exit Module: Captures vehicle images at the boom barrier, extracts the license plate using YOLO+LPRNet, and validates the entry.')
doc.add_paragraph('4. Dashboard & Analytics Module: Displays total capacity, currently available slots, and a searchable history of vehicle entries and exits.')

# 13
doc.add_heading('13. Work Plan and Timeline', level=2)
t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
h2 = t2.rows[0].cells
h2[0].text = 'Phase'
h2[1].text = 'Task Description'
h2[2].text = 'Duration'
h2[3].text = 'Week'

phases = [
    ('Phase 1', 'Requirement Analysis & Literature Survey', '2 Weeks', 'W1 - W2'),
    ('Phase 2', 'Dataset Collection & Preprocessing', '2 Weeks', 'W3 - W4'),
    ('Phase 3', 'YOLOv8 Model Training & Fine-Tuning', '3 Weeks', 'W5 - W7'),
    ('Phase 4', 'LPRNet Integration & Accuracy Testing', '2 Weeks', 'W8 - W9'),
    ('Phase 5', 'Backend API & Database Development', '2 Weeks', 'W10 - W11'),
    ('Phase 6', 'Frontend React Dashboard Development', '2 Weeks', 'W12 - W13'),
    ('Phase 7', 'System Integration, Testing & Bug Fixing', '2 Weeks', 'W14 - W15'),
    ('Phase 8', 'Final Documentation & Presentation Prep', '1 Week', 'W16')
]

for p, td, d, w in phases:
    r = t2.add_row().cells
    r[0].text = p
    r[1].text = td
    r[2].text = d
    r[3].text = w

doc.save('E:/projects/mtech/anti_smart parking/Review-I_Seminar_Presentation.docx')
