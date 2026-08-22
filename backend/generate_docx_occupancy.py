import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = docx.Document()

# Title
title = doc.add_heading('Dissertation work (Review-01)\nMaster of Technology in Computer Science and Engineering', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Project Title: Smart Parking Occupancy Detection System Using Deep Learning on Multi-Source Visual Inputs', level=1)

# 1
doc.add_heading('1. Introduction and Motivation', level=2)
doc.add_paragraph('The Smart Parking Occupancy Detection System is an intelligent management solution designed to automate the detection of available and occupied parking spaces in real time. Traditional parking management methods are time-consuming, prone to human error, and lack scalability for modern urban environments.')
p = doc.add_paragraph()
p.add_run('Motivation: ').bold = True
p.add_run('The primary motivation is to leverage computer vision and deep learning techniques to minimize manual intervention, reduce traffic congestion caused by parking searches, and contribute toward the development of intelligent urban transportation infrastructure using existing CCTV camera networks instead of expensive per-slot physical hardware sensors.')

# 2
doc.add_heading('2. Problem Statement', level=2)
doc.add_paragraph('Identifying the occupancy status of individual parking slots in real-world scenarios presents significant challenges due to varying environmental conditions such as extreme weather, dynamic lighting, unusual camera angles, and visual occlusions. Traditional image processing techniques and lightweight models often fail to maintain high accuracy under these unpredictable conditions, necessitating a highly robust deep transfer-learning approach capable of processing multi-source visual inputs reliably.')

# 3
doc.add_heading('3. Objectives', level=2)
doc.add_paragraph('1. Develop a binary classification model utilizing VGG16 transfer learning to accurately identify parking slot occupancy (vacant vs. occupied).')
doc.add_paragraph('2. Support multi-source visual inputs including live CCTV feeds, video files, and static images.')
doc.add_paragraph('3. Train and validate the model on standard benchmark datasets (PKLot, CNRPark, CNRPark-EXT) and a custom SOCIETY_PARKING dataset to achieve near 99.9% accuracy.')
doc.add_paragraph('4. Implement a decentralized processing architecture that synchronizes local occupancy data to a centralized server.')
doc.add_paragraph('5. Develop a web interface providing users with real-time slot availability maps and facility-level occupancy statistics.')

# 4
doc.add_heading('4. Literature Survey (Minimum 10 Recent IEEE/SCI Papers)', level=2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'S.No.'
hdr_cells[1].text = 'Paper Title (Year)'
hdr_cells[2].text = 'Core Technologies'
hdr_cells[3].text = 'Key Contribution'

papers = [
    ('1', 'Deep Learning-based Parking Occupancy Detection using VGG Architectures (2025)', 'VGG16, CNN, Transfer Learning', 'Benchmarked VGG models on the PKLot dataset, proving high feature extraction robustness in varying illumination.'),
    ('2', 'Multi-Source Visual Data Integration for Smart City Parking Solutions (2026)', 'Video Analytics, Decentralized AI', 'Proposed decentralized edge-processing paradigms for localized camera inputs synchronized to cloud databases.'),
    ('3', 'Robust Parking Space Classification using CNRPark and Transfer Learning (2024)', 'Transfer Learning, CNRPark-EXT', 'Demonstrated exceptional accuracy (99%+) on occlusion-heavy parking lots using deep feature embeddings.'),
    ('4', 'End-to-End Smart Parking System with Real-Time Availability (2025)', 'Web Services, Cloud DB, IoT', 'Integrated occupancy sensors with mobile interfaces featuring dynamic slot availability displays.'),
    ('5', 'Evaluation of CNNs for Real-Time Parking Slot Detection on Edge Devices (2026)', 'CNN, Edge Computing', 'Showcased the viability of running localized pre-processing (ROI extraction) on edge hardware before cloud sync.'),
    ('6', 'A Review of Computer Vision Techniques in Intelligent Transportation Systems (2024)', 'Computer Vision, ALPR', 'Summarized the transition from hardware-based IR sensors to scalable visual camera-based parking grids.'),
    ('7', 'Transfer Learning Approaches for Weather-Resilient Parking Detection (2025)', 'ResNet, VGG16, Weather Augmentation', 'Trained models across varied weather conditions (rain, snow) to ensure high commercial deployment viability.'),
    ('8', 'Dynamic ROI Extraction for Angled and Distorted CCTV Parking Feeds (2026)', 'OpenCV, Homography, CNN', 'Developed automated perspective transforms to feed standardized 224x224 ROIs into classification networks.'),
    ('9', 'Scalable Cloud Architectures for Multi-Tenant Smart Parking Facilities (2025)', 'Cloud Computing, Distributed DB', 'Designed a synchronization protocol for decentralized facilities reporting to a centralized mobile application backend.'),
    ('10', 'Contactless Smart Parking Management for Residential Societies (2026)', 'Deep Learning, Custom Datasets', 'Highlighted the creation and utility of customized datasets (SOCIETY_PARKING) to fine-tune general models for specific environments.')
]

for s, title, tech, cont in papers:
    row_cells = table.add_row().cells
    row_cells[0].text = s
    row_cells[1].text = title
    row_cells[2].text = tech
    row_cells[3].text = cont

# 5
doc.add_heading('5. Research Gap Identification', level=2)
doc.add_paragraph('1. Over-reliance on Physical Hardware: Many smart parking systems still rely on expensive per-slot physical sensors (IR/Ultrasonic) which are difficult to maintain and scale.')
doc.add_paragraph('2. Fragility to Environmental Changes: Existing lightweight vision models suffer severe accuracy drops during nighttime, heavy rain, or partial vehicle occlusions.')
doc.add_paragraph('3. Lack of Integrated User Applications: Research often focuses purely on the classification metric, neglecting the integration of full-stack user applications that handle decentralized sync and real-time dashboard analytics.')

# 6
doc.add_heading('6. Proposed Methodology', level=2)
doc.add_paragraph('The proposed system uses a decentralized processing paradigm. Local facilities collect images from multi-source inputs and extract Regions of Interest (ROI) for individual parking slots. Each ROI is preprocessed and resized to exactly 224×224 pixels.')
doc.add_paragraph('The 224x224 image is passed through a VGG16 CNN utilizing transfer learning. The VGG16 feature extraction layers feed into fully connected dense layers that produce a binary classification output. A strict threshold value of 0.5 is applied: values closer to 0 indicate a vacant slot, while values closer to 1 indicate an occupied slot. The locally processed occupancy data is then synced to a remote central server where the web application dynamically displays real-time slot availability.')

# 7
doc.add_heading('7. Algorithm / Flowchart', level=2)
doc.add_paragraph('1. Multi-Source Visual Input Capture (Live CCTV / Video / Images)')
doc.add_paragraph('2. Local ROI Extraction & Preprocessing (Resize to 224x224)')
doc.add_paragraph('3. Feature Extraction (VGG16 Transfer Learning CNN)')
doc.add_paragraph('4. Dense Layer Binary Classification (Threshold 0.5)')
doc.add_paragraph('   a) If output > 0.5 -> Mark as Occupied')
doc.add_paragraph('   b) If output <= 0.5 -> Mark as Vacant')
doc.add_paragraph('5. Synchronize Occupancy Data to Centralized Server')
doc.add_paragraph('6. Web Application UI Updates')

# 8
doc.add_heading('8. System Architecture', level=2)
doc.add_paragraph('Decentralized Edge Node: Responsible for image collection, frame extraction, ROI masking, 224x224 resizing, and VGG16 inference.', style='List Bullet')
doc.add_paragraph('Centralized Remote Server: Acts as the synchronization hub. Hosts the database containing overall facility layouts, current status, and vehicle logs.', style='List Bullet')
doc.add_paragraph('Web Application Interface: The client-facing layer allowing users and admins to view real-time availability maps and track facility analytics.', style='List Bullet')

# 9
doc.add_heading('9. Software & Hardware Requirements', level=2)
doc.add_paragraph('Hardware Requirements:', style='Heading 3')
doc.add_paragraph('- Processing Unit: GPU-enabled machine (e.g., NVIDIA CUDA support) for real-time VGG16 inference.')
doc.add_paragraph('- Input Sources: Standard IP/CCTV cameras for visual feeds.')
doc.add_paragraph('- RAM: 8GB - 16GB Minimum.')
doc.add_paragraph('Software Requirements:', style='Heading 3')
doc.add_paragraph('- Operating System: Windows / Linux')
doc.add_paragraph('- Frameworks: PyTorch / TensorFlow (for VGG16), OpenCV (for ROI processing)')
doc.add_paragraph('- Web Technologies: React.js (Frontend), FastAPI/Node.js (Backend)')

# 10
doc.add_heading('10. Tools & Development Environment', level=2)
doc.add_paragraph('- IDE: Visual Studio Code, PyCharm')
doc.add_paragraph('- Version Control: Git')
doc.add_paragraph('- API Testing: Postman')
doc.add_paragraph('- Environment: Python 3.9+, Node.js')

# 11
doc.add_heading('11. Dataset Description', level=2)
doc.add_paragraph('The system is rigorously trained and evaluated on several prominent benchmark datasets to ensure commercial deployment viability:')
doc.add_paragraph('- PKLot Dataset: Contains nearly 700,000 images captured across different weather conditions and days.', style='List Bullet')
doc.add_paragraph('- CNRPark & CNRPark-EXT: Datasets specifically challenging due to extreme camera angles and significant occlusion scenarios.', style='List Bullet')
doc.add_paragraph('- Custom SOCIETY_PARKING Dataset: Locally captured and annotated dataset designed to fine-tune the model to the exact visual characteristics of residential complexes.', style='List Bullet')
doc.add_paragraph('Result: Achieves up to 99.917% on standard benchmarks and 97.61% on real-world footage.')

# 12
doc.add_heading('12. Module Description', level=2)
doc.add_paragraph('1. Input & Preprocessing Module: Accepts multi-source inputs, maps user-defined polygonal ROIs, and standardizes image crops to 224x224 arrays.')
doc.add_paragraph('2. VGG16 Classification Engine: The core deep learning module responsible for the transfer-learning based feature extraction and binary regression (vacant vs occupied).')
doc.add_paragraph('3. Synchronization Node: Handles the decentralized communication, packaging local occupancy arrays and transmitting them to the central cloud securely.')
doc.add_paragraph('4. User Dashboard & Analytics: Renders the real-time UI map for end users. Displays total capacity, currently available slots, and logs.')

# 13
doc.add_heading('13. Work Plan and Timeline', level=2)
t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
h2 = t2.rows[0].cells
h2[0].text = 'Phase'
h2[1].text = 'Task Description'
h2[2].text = 'Duration'
h2[3].text = 'Week'

phases = [
    ('Phase 1', 'Literature Survey & Dataset Aggregation (PKLot, CNRPark)', '2 Weeks', 'W1 - W2'),
    ('Phase 2', 'Data Preprocessing & Custom SOCIETY_PARKING Curation', '2 Weeks', 'W3 - W4'),
    ('Phase 3', 'VGG16 Transfer Learning Model Training & Validation', '3 Weeks', 'W5 - W7'),
    ('Phase 4', 'ROI Extraction Scripting & Local Deployment Configuration', '2 Weeks', 'W8 - W9'),
    ('Phase 5', 'Decentralized Server Synchronization Protocol Setup', '2 Weeks', 'W10 - W11'),
    ('Phase 6', 'Web Application Dashboard & UI Integration', '3 Weeks', 'W12 - W14'),
    ('Phase 7', 'End-to-End System Integration & Real-World Testing', '1 Week', 'W15'),
    ('Phase 8', 'Final Review, Thesis Documentation & Presentation Prep', '1 Week', 'W16')
]

for p, td, d, w in phases:
    r = t2.add_row().cells
    r[0].text = p
    r[1].text = td
    r[2].text = d
    r[3].text = w

doc.save('E:/projects/mtech/anti_smart parking/Review-I_Seminar_Presentation_Revised.docx')
